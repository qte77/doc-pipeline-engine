# Copyright 2026 qte77
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
"""Render-formats utility: Markdown → MD + DOCX + PDF artifacts.

Both V1 (Claude-written Markdown) and V2 (Jinja-rendered Markdown) feed
through the same converter so the A/B between variants is purely about
the Markdown content, not the conversion path.
"""
from __future__ import annotations

import io

import pytest

pytest.importorskip("docx")
pytest.importorskip("markdown")
pytest.importorskip("weasyprint")

from doc_pipeline_engine.render.formats import RenderArtifacts, render_artifacts


def test_render_formats_md_is_input_passthrough() -> None:
    md = "# Title\n\nBody text."

    art = render_artifacts(md, title="Test")

    assert art.md == md


def test_render_formats_returns_render_artifacts_with_three_fields() -> None:
    art = render_artifacts("# Hello", title="Test")

    assert isinstance(art, RenderArtifacts)
    assert isinstance(art.md, str)
    assert isinstance(art.docx, bytes)
    assert isinstance(art.pdf, bytes)


def test_render_formats_pdf_starts_with_pdf_magic() -> None:
    art = render_artifacts("# Hello\n\nWorld.", title="Test")

    assert art.pdf[:5] == b"%PDF-"


def test_render_formats_docx_is_valid_ooxml_zip() -> None:
    art = render_artifacts("# Hello", title="Test")

    # OOXML files are zip archives — magic bytes PK\x03\x04
    assert art.docx[:4] == b"PK\x03\x04"


def test_render_formats_docx_round_trips_through_python_docx() -> None:
    from docx import Document

    art = render_artifacts("# Hello\n\nBody.", title="Test")

    doc = Document(io.BytesIO(art.docx))
    texts = [p.text for p in doc.paragraphs]
    assert "Hello" in texts
    assert "Body." in texts


def test_render_formats_headings_become_docx_heading_paragraphs() -> None:
    from docx import Document

    art = render_artifacts("# Top\n\n## Sub\n\nBody.", title="Test")

    doc = Document(io.BytesIO(art.docx))
    styles_by_text = {p.text: p.style.name for p in doc.paragraphs}
    assert styles_by_text["Top"].startswith("Heading")
    assert styles_by_text["Sub"].startswith("Heading")
    assert styles_by_text["Body."].startswith("Normal")
